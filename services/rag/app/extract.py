"""Getting text out of the formats people actually have.

This is the first thing the split buys. The TypeScript upload accepted plain
text only, because parsing PDF and DOCX in Node means either a weak pure-JS
parser or a native dependency that fights the serverless bundle. Here it is two
mature libraries.

Extraction is where a RAG pipeline quietly loses: a PDF parsed without regard
for layout yields interleaved column text, and no amount of chunking or
reranking recovers meaning that was scrambled before it was ever stored.
PyMuPDF's "blocks" mode preserves reading order, which is why it is used
instead of the simpler whole-page `get_text()`.

Both parsers are imported lazily so a deployment that only ingests Markdown
does not need the `extract` extra installed.
"""

from __future__ import annotations

import io


class ExtractionError(Exception):
    """Raised when a document cannot be turned into text."""


def _extract_pdf(data: bytes) -> str:
    try:
        import pymupdf
    except ImportError as error:  # pragma: no cover - depends on install extras
        raise ExtractionError(
            "PDF support needs the 'extract' extra: pip install '.[extract]'"
        ) from error

    try:
        document = pymupdf.open(stream=data, filetype="pdf")
    except Exception as error:
        raise ExtractionError(f"Could not open the PDF: {error}") from error

    pages: list[str] = []
    with document:
        for number, page in enumerate(document, start=1):
            # sort=True orders blocks by position rather than by the order they
            # happen to appear in the content stream, which is what keeps a
            # two-column layout from interleaving line by line.
            text = page.get_text("text", sort=True).strip()
            if not text:
                continue
            # Page numbers become headings, so a citation can point at a page.
            pages.append(f"## Page {number}\n\n{text}")

    if not pages:
        raise ExtractionError(
            "No text found. This looks like a scanned PDF — it needs OCR first."
        )
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as error:  # pragma: no cover - depends on install extras
        raise ExtractionError(
            "DOCX support needs the 'extract' extra: pip install '.[extract]'"
        ) from error

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as error:
        raise ExtractionError(f"Could not open the document: {error}") from error

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # Word's built-in heading styles map onto Markdown headings, which the
        # chunker then uses as split points and heading trails. Structure the
        # author already applied is structure retrieval gets for free.
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    if not parts:
        raise ExtractionError("The document has no readable text.")
    return "\n\n".join(parts)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 maps every byte, so reaching here means the loop above changed.
    raise ExtractionError("Could not decode the file as text.")


# Extensions are checked before MIME types: browsers report .md as everything
# from text/markdown to application/octet-stream depending on the platform.
_PDF = {".pdf"}
_DOCX = {".docx"}


def extract_text(data: bytes, filename: str = "", mime_type: str = "") -> str:
    """Turn an uploaded file into Markdown-ish plain text."""
    lower = filename.lower()

    if any(lower.endswith(ext) for ext in _PDF) or mime_type == "application/pdf":
        return _extract_pdf(data)

    if any(lower.endswith(ext) for ext in _DOCX) or mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(data)

    if lower.endswith(".doc"):
        raise ExtractionError(
            "Legacy .doc files are not supported. Save it as .docx or PDF first."
        )

    return _decode_text(data)
